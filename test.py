from docx import Document
from docx.shared import Pt
from charset_normalizer import from_path
import re

def detect_encoding(file_path):
    result = from_path(file_path).best()
    if result is None:
        print("Không phát hiện được encoding, dùng utf-8 mặc định.")
        return 'utf-8'
    return result.encoding

def clean_text(text: str) -> str:
    """Loại bỏ ký tự control không hợp lệ cho DOCX."""
    return re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

def set_run_font(run, font_name='Times New Roman', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)

def txt_to_docx(input_path: str, output_path: str):
    doc = Document()
    encoding = detect_encoding(input_path)
    print(f"Detected encoding: {encoding}")
    
    with open(input_path, 'r', encoding=encoding, errors="ignore") as f:
        lines = f.readlines()
    
    for line in lines:
        stripped = clean_text(line.strip())
        if not stripped:
            continue
        
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(stripped)
        
        # Set font cho paragraph/heading
        for run in p.runs:
            set_run_font(run, font_name='Times New Roman', font_size=12)
    
    doc.save(output_path)
    print(f"Đã chuyển {input_path} sang {output_path}")

if __name__ == "__main__":
    txt_file = "rag_output.docx"
    docx_file = "de_thi.docx"
    txt_to_docx(txt_file, docx_file)
